from __future__ import annotations

from pathlib import Path

from app.api.schemas import ParsedResume, TailoredResume
from app.core.logging import get_logger

logger = get_logger(__name__)


def generate_ats_docx(
    resume: ParsedResume,
    tailored: TailoredResume,
    output_path: Path,
) -> Path:
    """Generate an ATS-friendly DOCX resume.

    Format rules:
    - Single column, no tables/graphics
    - Standard headings
    - Calibri font, 11pt body, 14pt headings
    - Reverse-chronological experience
    - Standard bullet points
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Header: Name ---
    name = resume.candidate_name or "Candidate"
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(18)
        run.font.name = "Calibri"

    # --- Contact line ---
    contact_parts = []
    if resume.email:
        contact_parts.append(resume.email)
    if resume.phone:
        contact_parts.append(resume.phone)
    if resume.location:
        contact_parts.append(resume.location)
    if resume.linkedin:
        contact_parts.append(resume.linkedin)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.style.font.size = Pt(10)

    # --- Professional Summary ---
    if tailored.tailored_summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(tailored.tailored_summary)

    # --- Skills ---
    if tailored.tailored_skills:
        doc.add_heading("Skills", level=1)
        # Group into lines of ~8 skills
        skills = tailored.tailored_skills
        for i in range(0, len(skills), 8):
            chunk = skills[i:i+8]
            doc.add_paragraph(" | ".join(chunk))

    # --- Experience ---
    if tailored.tailored_experience:
        doc.add_heading("Professional Experience", level=1)
        for exp in tailored.tailored_experience:
            # Job title + company
            title_line = exp.job_title or "Role"
            if exp.company:
                title_line += f" — {exp.company}"
            title_para = doc.add_paragraph()
            run = title_para.add_run(title_line)
            run.bold = True
            run.font.size = Pt(11)

            # Dates
            if exp.start_date:
                date_str = f"{exp.start_date} – {exp.end_date or 'Present'}"
                date_para = doc.add_paragraph(date_str)
                date_para.runs[0].italic = True
                date_para.runs[0].font.size = Pt(10)

            # Bullets
            for bullet in (exp.description or [])[:6]:
                doc.add_paragraph(bullet, style="List Bullet")

    # --- Education ---
    if resume.education:
        doc.add_heading("Education", level=1)
        for edu in resume.education:
            edu_line = edu.degree or "Degree"
            if edu.field_of_study:
                edu_line += f" in {edu.field_of_study}"
            if edu.institution:
                edu_line += f" — {edu.institution}"
            if edu.graduation_date:
                edu_line += f" ({edu.graduation_date})"
            doc.add_paragraph(edu_line)

    # --- Certifications ---
    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in resume.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Generated ATS DOCX: %s", output_path)
    return output_path
