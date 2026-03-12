"""Generate cover letter DOCX files in ATS-friendly format."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from app.api.schemas import ParsedResume
from app.core.logging import get_logger

logger = get_logger(__name__)


def generate_cover_letter_docx(
    text: str,
    resume: ParsedResume,
    output_path: Path,
) -> Path:
    """Generate an ATS-friendly cover letter DOCX.

    Formal letter format with date, greeting, body paragraphs, sign-off.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header: candidate contact info
    contact_parts = []
    if resume.candidate_name:
        name_para = doc.add_paragraph()
        run = name_para.add_run(resume.candidate_name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Calibri"

    contact_items = []
    if resume.email:
        contact_items.append(resume.email)
    if resume.phone:
        contact_items.append(resume.phone)
    if resume.location:
        contact_items.append(resume.location)
    if contact_items:
        doc.add_paragraph(" | ".join(contact_items))

    # Date
    doc.add_paragraph("")  # spacer
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph("")  # spacer

    # Body paragraphs
    for paragraph in text.strip().split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        # Check if it's a sign-off line (short, like "Sincerely,")
        if paragraph.startswith("Sincerely") or paragraph.startswith("Best regards"):
            doc.add_paragraph("")
            doc.add_paragraph(paragraph)
        elif paragraph == resume.candidate_name:
            name_para = doc.add_paragraph()
            run = name_para.add_run(paragraph)
            run.bold = True
        else:
            doc.add_paragraph(paragraph)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Generated cover letter DOCX: %s", output_path)
    return output_path
